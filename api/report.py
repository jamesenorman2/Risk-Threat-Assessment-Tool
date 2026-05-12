"""Word report generator (Vercel Python serverless).

POST /api/report with an assessment JSON body (same shape as a saved project
file in projects/). Returns a .docx download.

Report order: cover -> executive summary -> section 6 (Security Risk Assessment)
  6.1 Assessment Methodology
  6.2 Asset Identification
  6.3 Asset Criticality
  6.4 Threats
    6.4.1 Common Threat Actors in the UK
    6.4.2 Threat Register
    6.4.3 Threat-Asset Mapping
  6.5 Attractiveness, Vulnerability & Likelihood
  6.6 Points of Interest
  6.7 Impact Assessment
  6.8 Existing Controls
  6.9 Risk Register (Pre Mitigation)
  6.10 Security Treatments
"""
from http.server import BaseHTTPRequestHandler
import io
import json
import os
import re
from datetime import datetime

os.environ.setdefault("MPLCONFIGDIR", "/tmp/mpl")
os.makedirs(os.environ["MPLCONFIGDIR"], exist_ok=True)

def _mpl():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as _plt
    from matplotlib.patches import Rectangle as _Rect
    return _plt, _Rect

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GREY_DARK = RGBColor(0x47, 0x55, 0x69)
GREY_MID = RGBColor(0x64, 0x74, 0x8B)
GREY_LIGHT = RGBColor(0xD6, 0xDC, 0xE4)

LEVEL_HEX = {
    "Very Low": "4472C4",
    "Low": "70AD47",
    "Medium": "FFC000",
    "High": "ED7D31",
    "Very High": "C00000",
}
LEVEL_FG = {
    "Very Low": RGBColor(0xFF, 0xFF, 0xFF),
    "Low": RGBColor(0xFF, 0xFF, 0xFF),
    "Medium": RGBColor(0x1A, 0x1A, 0x1A),
    "High": RGBColor(0xFF, 0xFF, 0xFF),
    "Very High": RGBColor(0xFF, 0xFF, 0xFF),
}
SCORE_LEVEL = {1: "Very Low", 2: "Low", 3: "Medium", 4: "High", 5: "Very High"}
IL = {1: "Negligible", 2: "Minor", 3: "Moderate", 4: "Severe", 5: "Catastrophic"}
LL = {1: "Very Unlikely", 2: "Unlikely", 3: "Possible", 4: "Likely", 5: "Very Likely"}
RM = [
    ["Medium", "High", "High", "Very High", "Very High"],
    ["Low", "Medium", "High", "High", "Very High"],
    ["Low", "Medium", "Medium", "High", "High"],
    ["Very Low", "Low", "Medium", "Medium", "High"],
    ["Very Low", "Very Low", "Low", "Low", "Medium"],
]
TAC = ["visibility", "status", "access", "collateral", "crime", "history"]
TAC_LABELS = {
    "visibility": "Visibility", "status": "Status", "access": "Access",
    "collateral": "Collateral", "crime": "Crime Profile", "history": "Historical",
}
IC = ["disruption", "fatalities", "reputation", "financial", "property"]
IC_LABELS = {
    "disruption": "Disruption", "fatalities": "Fatalities", "reputation": "Reputation",
    "financial": "Financial", "property": "Property",
}
AC_KEYS = ["opImpact", "dependencies", "impactPeople", "impactEnv"]
AC_LABELS = {
    "opImpact": "Op. Impact", "dependencies": "Dependencies",
    "impactPeople": "People", "impactEnv": "Environment",
}

CRIT_HEX = {
    "Critical": "C00000", "High": "ED7D31", "Medium": "FFC000",
    "Low": "70AD47", "Negligible": "4472C4",
}
CRIT_FG = {
    "Critical": RGBColor(0xFF, 0xFF, 0xFF), "High": RGBColor(0xFF, 0xFF, 0xFF),
    "Medium": RGBColor(0x1A, 0x1A, 0x1A), "Low": RGBColor(0xFF, 0xFF, 0xFF),
    "Negligible": RGBColor(0xFF, 0xFF, 0xFF),
}

GAP_HEX = {"Adequate": "E2EFDA", "Partial": "FFF2CC", "Insufficient": "FFC7CE"}
GAP_FG = {
    "Adequate": RGBColor(0x16, 0xA3, 0x4A),
    "Partial": RGBColor(0xD9, 0x77, 0x06),
    "Insufficient": RGBColor(0xDC, 0x26, 0x26),
}


# --- Computation helpers ---

def _avg(values):
    vals = [v for v in values if v is not None and v != 0 and v != ""]
    if not vals:
        return None
    return round(sum(vals) / len(vals), 1)


def _get_rr(lik, imp):
    if not lik or not imp:
        return None
    lik_i = int(round(lik))
    imp_i = int(round(imp))
    if not (1 <= lik_i <= 5 and 1 <= imp_i <= 5):
        return None
    return RM[5 - lik_i][imp_i - 1]


def _asset_rank(score):
    if score is None:
        return ""
    if score >= 4.5:
        return "Critical"
    if score >= 3.5:
        return "High"
    if score >= 2.5:
        return "Medium"
    if score >= 1.5:
        return "Low"
    return "Negligible"


def compute_risks(state):
    threats = state.get("threats", []) or []
    vs = state.get("vs", {}) or {}
    is_ = state.get("is_", {}) or {}
    tx = state.get("tx", {}) or {}
    rs = state.get("rs", {}) or {}
    acc = state.get("acc", {}) or {}

    risks = []
    for t in threats:
        tid = str(t.get("id"))
        cap, intent = t.get("cap"), t.get("int")
        sc = round((cap + intent) / 2) if cap and intent else None
        lv = SCORE_LEVEL.get(sc) if sc else None
        dbt = bool(sc and sc >= 3)

        v = vs.get(tid, {}) or {}
        im = is_.get(tid, {}) or {}
        ta = _avg([v.get(c) for c in TAC])
        vu = v.get("vuln")
        tot = sc + ta + vu if (sc and ta and vu) else None
        lik = round(tot / 3) if tot else None
        raw_imp = _avg([im.get(c) for c in IC])
        fat = im.get("fatalities")
        life_safety = bool(fat and fat >= 4 and raw_imp and raw_imp < 4)
        imp = 4 if life_safety else raw_imp
        rat = _get_rr(lik, imp)

        r = rs.get(tid, {}) or {}
        res_l, res_i = r.get("l"), r.get("i")
        res_rat = _get_rr(res_l, res_i)

        risks.append({
            "id": t.get("id"),
            "n": t.get("n", ""),
            "d": t.get("d", ""),
            "cat": t.get("cat", ""),
            "cap": cap, "int": intent, "sc": sc, "lv": lv, "dbt": dbt,
            "ta": ta, "vu": vu, "lik": lik, "imp": imp,
            "rawImp": raw_imp, "lifeSafetyOverride": life_safety,
            "rat": rat,
            "treat": tx.get(tid),
            "resL": res_l, "resI": res_i, "resRat": res_rat,
            "acc": acc.get(tid),
            "targets": t.get("targets", []) or [],
        })
    return risks


# --- Low-level docx helpers ---

def _shade(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _borders(table, size="4", color="D6DCE4"):
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tbl_pr.append(borders)


def _set_cell_text(cell, text, *, bold=False, size=9, color=None, align="left",
                   fill=None, font="Calibri"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(str(text) if text is not None else "")
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if fill:
        _shade(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = NAVY
        _set_para_border(p, bottom=("12", "1B2A4A"))
        _set_para_spacing(p, before=180, after=120)
    elif level == 2:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        _set_para_spacing(p, before=120, after=80)
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = GREY_DARK
        _set_para_spacing(p, before=80, after=40)
    return p


def _set_para_spacing(p, before=0, after=0):
    pf = p.paragraph_format
    if before:
        pf.space_before = Pt(before / 20)
    if after:
        pf.space_after = Pt(after / 20)


def _set_para_border(p, bottom=None):
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    if bottom:
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), bottom[0])
        b.set(qn("w:color"), bottom[1])
        b.set(qn("w:space"), "1")
        pbdr.append(b)
    p_pr.append(pbdr)


def _add_para(doc, text, *, size=10, bold=False, italic=False, color=None,
              align="left", space_after=60):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    run = p.add_run(text or "")
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    _set_para_spacing(p, after=space_after)
    return p


def _page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _hdr_row(tbl, headers, widths_cm):
    for col, w in zip(tbl.columns, widths_cm):
        col.width = Cm(w)
    for i, h in enumerate(headers):
        _set_cell_text(tbl.rows[0].cells[i], h, bold=True, size=8,
                       color=RGBColor(0xFF, 0xFF, 0xFF), fill="1B2A4A",
                       align="center")


# --- Chart helpers ---

def _try_picture(doc, builder, width):
    try:
        buf = builder()
        doc.add_picture(buf, width=width)
    except Exception as e:
        p = doc.add_paragraph()
        run = p.add_run(f"[Chart unavailable: {type(e).__name__}]")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.italic = True
        run.font.color.rgb = GREY_MID


def _hex_to_rgb(h):
    return tuple(int(h[i:i + 2], 16) / 255 for i in (0, 2, 4))


def chart_risk_matrix(risks, *, residual=False, title=""):
    plt, Rectangle = _mpl()
    fig, ax = plt.subplots(figsize=(6.5, 4.8), dpi=150)
    for i in range(5):
        for j in range(5):
            level = RM[i][j]
            ax.add_patch(Rectangle(
                (j, 4 - i), 1, 1,
                facecolor=_hex_to_rgb(LEVEL_HEX[level]),
                edgecolor="white", linewidth=2,
            ))
    for idx, r in enumerate(risks):
        lik = r.get("resL") if residual else r.get("lik")
        imp = r.get("resI") if residual else r.get("imp")
        if not lik or not imp:
            continue
        lik_i = int(round(lik)); imp_i = int(round(imp))
        if not (1 <= lik_i <= 5 and 1 <= imp_i <= 5):
            continue
        ax.text(imp_i - 0.5, lik_i - 0.5, f"R{idx + 1}",
                ha="center", va="center", fontsize=8, fontweight="bold",
                color="white" if RM[5 - lik_i][imp_i - 1] in ("Very Low", "Low", "High", "Very High") else "black")
    ax.set_xlim(0, 5); ax.set_ylim(0, 5)
    ax.set_xticks([i + 0.5 for i in range(5)])
    ax.set_yticks([i + 0.5 for i in range(5)])
    ax.set_xticklabels([f"{i}\n{IL[i]}" for i in range(1, 6)], fontsize=7)
    ax.set_yticklabels([f"{i}\n{LL[i]}" for i in range(1, 6)], fontsize=7)
    ax.set_xlabel("Impact", fontsize=9, fontweight="bold")
    ax.set_ylabel("Likelihood", fontsize=9, fontweight="bold")
    if title:
        ax.set_title(title, fontsize=11, fontweight="bold", color="#1B2A4A", pad=10)
    ax.tick_params(left=False, bottom=False)
    for spine in ax.spines.values():
        spine.set_visible(False)
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def chart_profile(risks, *, residual=False, title=""):
    plt, _ = _mpl()
    order = ["Very Low", "Low", "Medium", "High", "Very High"]
    key = "resRat" if residual else "rat"
    counts = [sum(1 for r in risks if r.get(key) == lv) for lv in order]
    fig, ax = plt.subplots(figsize=(6.5, 3.0), dpi=150)
    bars = ax.bar(order, counts,
                  color=[_hex_to_rgb(LEVEL_HEX[lv]) for lv in order],
                  edgecolor="white", linewidth=1.2)
    for bar, c in zip(bars, counts):
        if c:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                    str(c), ha="center", va="bottom", fontsize=10, fontweight="bold")
    ax.set_ylabel("Number of risks", fontsize=9)
    ax.set_ylim(0, max(counts + [1]) * 1.25)
    ax.tick_params(axis="x", labelsize=9)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    if title:
        ax.set_title(title, fontsize=11, fontweight="bold", color="#1B2A4A", pad=8)
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def chart_pre_post(risks):
    plt, _ = _mpl()
    order = ["Very Low", "Low", "Medium", "High", "Very High"]
    pre = [sum(1 for r in risks if r.get("rat") == lv) for lv in order]
    post = [sum(1 for r in risks if r.get("resRat") == lv) for lv in order]
    x = list(range(len(order)))
    w = 0.38
    fig, ax = plt.subplots(figsize=(6.5, 3.2), dpi=150)
    b1 = ax.bar([i - w / 2 for i in x], pre, w, label="Pre-treatment",
                color=[_hex_to_rgb(LEVEL_HEX[lv]) for lv in order],
                edgecolor="#333", linewidth=0.8)
    b2 = ax.bar([i + w / 2 for i in x], post, w, label="Residual",
                color=[_hex_to_rgb(LEVEL_HEX[lv]) for lv in order],
                edgecolor="#333", linewidth=0.8, hatch="///", alpha=0.75)
    for bars in (b1, b2):
        for bar in bars:
            h = bar.get_height()
            if h:
                ax.text(bar.get_x() + bar.get_width() / 2, h + 0.1, str(int(h)),
                        ha="center", va="bottom", fontsize=8, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(order, fontsize=9)
    ax.set_ylabel("Number of risks", fontsize=9)
    ax.set_ylim(0, max(pre + post + [1]) * 1.3)
    ax.legend(loc="upper right", fontsize=9, frameon=False)
    ax.set_title("Pre-Treatment vs Residual Risk Profile",
                 fontsize=11, fontweight="bold", color="#1B2A4A", pad=8)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    buf = io.BytesIO()
    plt.tight_layout()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# --- Cover & Executive Summary ---

def _build_cover(doc, ctx, risks):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SECURITY RISK & THREAT ASSESSMENT")
    run.font.name = "Calibri"; run.font.size = Pt(10); run.font.bold = True
    run.font.color.rgb = GREY_MID
    _set_para_spacing(p, after=120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(ctx.get("project") or "Assessment")
    run.font.name = "Calibri"; run.font.size = Pt(28); run.font.bold = True
    run.font.color.rgb = NAVY
    _set_para_spacing(p, after=80)

    addr = " ".join([x for x in (ctx.get("address"), ctx.get("postcode")) if x])
    if addr:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(addr)
        run.font.name = "Calibri"; run.font.size = Pt(12); run.font.color.rgb = GREY_DARK
        _set_para_spacing(p, after=200)

    dbts = [r for r in risks if r.get("dbt")]
    scored = [r for r in risks if r.get("sc") is not None]

    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _borders(tbl, color="D6DCE4")
    tbl.columns[0].width = Cm(8); tbl.columns[1].width = Cm(8)

    left_lines = [
        ("Project No", ctx.get("projectNum") or "—"),
        ("Reference", ctx.get("ref") or "—"),
        ("Date", ctx.get("date") or "—"),
        ("Stage", ctx.get("stage") or "—"),
        ("Development", ctx.get("devType") or "—"),
    ]
    right_lines = [
        ("Assessor", ctx.get("assessor") or "—"),
        ("Client", ctx.get("client") or "—"),
        ("Threats Assessed", str(len(scored))),
        ("Design Basis Threats", str(len(dbts))),
    ]
    _set_cell_text(tbl.rows[0].cells[0], "Project Details", bold=True, size=10,
                   color=NAVY, fill="F1F5F9")
    _set_cell_text(tbl.rows[0].cells[1], "Assessment Details", bold=True, size=10,
                   color=NAVY, fill="F1F5F9")
    _set_cell_text(tbl.rows[1].cells[0], "", fill="FAFBFC")
    _set_cell_text(tbl.rows[1].cells[1], "", fill="FAFBFC")

    def _fill_block(cell, lines):
        cell.text = ""
        for label, val in lines:
            p = cell.add_paragraph()
            r = p.add_run(f"{label}: "); r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = NAVY
            r2 = p.add_run(str(val)); r2.font.size = Pt(9); r2.font.color.rgb = GREY_DARK
            _set_para_spacing(p, after=40)

    _fill_block(tbl.rows[1].cells[0], left_lines)
    _fill_block(tbl.rows[1].cells[1], right_lines)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated {datetime.utcnow().strftime('%d %B %Y')}  |  v{ctx.get('_v') or ''}")
    run.font.name = "Calibri"; run.font.size = Pt(8); run.font.color.rgb = GREY_MID

    _page_break(doc)


def _build_exec_summary(doc, ctx, risks, pois):
    _add_heading(doc, "Executive Summary", 1)

    scored = [r for r in risks if r.get("sc") is not None]
    dbts = [r for r in risks if r.get("dbt")]
    terror = [r for r in risks if r.get("cat") == "Terrorism"]
    crime = [r for r in risks if r.get("cat") and r.get("cat") != "Terrorism"]
    headline = [r for r in risks if r.get("rat") in ("Very High", "High")]

    intro = (
        f"This Security Risk Assessment evaluates {len(scored)} threats "
        f"({len(terror)} terrorism, {len(crime)} crime) identified for "
        f"{ctx.get('project') or 'the project'}. "
        f"{len(dbts)} threats meet the Design Basis Threat (DBT) criterion "
        f"(score ≥ 3), and {len(headline)} carry a pre-treatment rating of High or Very High. "
        f"{len(pois)} Points of Interest have been mapped against applicable threats."
    )
    _add_para(doc, intro, size=10, align="justify", space_after=120)

    _add_heading(doc, "Risk Profile Overview", 2)
    _try_picture(doc, lambda: chart_profile(risks, residual=False,
                                            title="Pre-Treatment Risk Profile"), Cm(16))

    reduced = sum(1 for r in risks
                  if r.get("rat") and r.get("resRat") and
                  ["Very Low", "Low", "Medium", "High", "Very High"].index(r["resRat"]) <
                  ["Very Low", "Low", "Medium", "High", "Very High"].index(r["rat"]))
    _add_para(doc,
              f"Following the proposed treatment strategy, {reduced} of "
              f"{len([r for r in risks if r.get('resRat')])} assessed threats show a "
              f"reduction in risk rating. Residual risk distribution is shown below.",
              size=10, align="justify", space_after=120)
    _try_picture(doc, lambda: chart_pre_post(risks), Cm(16))

    _page_break(doc)


# --- Section 6: Security Risk Assessment ---

def _build_s6_header(doc):
    _add_heading(doc, "6  Security Risk Assessment", 1)


def _build_methodology(doc):
    _add_heading(doc, "6.1  Assessment Methodology", 2)
    _add_para(doc,
              "This Security Risk Assessment has been conducted using the established SRA methodology, "
              "developed in full alignment with internationally recognised risk management standards. "
              "The methodology provides a structured, evidence-based framework following a logical sequence: "
              "establishing the security context and stakeholder consultation; asset identification and "
              "criticality ranking; threat identification categorised into Terrorism and Crime; Points of "
              "Interest mapping; threat analysis based on capability and intent; target attractiveness and "
              "vulnerability assessment; likelihood and impact scoring; risk rating via the 5×5 matrix; "
              "control effectiveness evaluation; risk treatment; and residual risk assessment.",
              size=10, align="justify", space_after=100)

    _add_heading(doc, "ISO 31000:2018 — Risk Management Guidelines", 3)
    _add_para(doc,
              "The SRA implements ISO 31000:2018: Scope, Context and Criteria (Clause 6.3); Risk Assessment "
              "comprising identification, analysis and evaluation (Clause 6.4); Risk Treatment via the TAAR "
              "framework (Clause 6.5); Communication and Consultation (Clause 6.2); and Monitoring and "
              "Review (Clause 6.6).",
              size=10, align="justify", space_after=80)

    _add_heading(doc, "AS/NZS HB 167:2006 — Security Risk Management", 3)
    _add_para(doc,
              "Structured around HB167: critical asset identification and ranking; dual-axis "
              "capability/intent threat assessment and DBT filtering; target attractiveness across six "
              "criteria; vulnerability assessment; likelihood determination; consequence analysis; existing "
              "control evaluation; and residual risk post-treatment.",
              size=10, align="justify", space_after=80)

    _add_heading(doc, "ISO/IEC 31010:2019 — Risk Assessment Techniques", 3)
    _add_para(doc,
              "Primary technique: semi-quantitative consequence/likelihood matrix (Technique B.29), "
              "supplemented by structured threat/scenario analysis (B.2) and expert judgement (B.1). "
              "Known limitations are acknowledged and mitigated through structured scoring descriptors "
              "and iterative reassessment.",
              size=10, align="justify", space_after=80)

    _add_heading(doc, "Calculation Methodology", 3)

    rows = [
        ("Threat Score", "ROUND((Capability + Intent) / 2). Design Basis Threats identified at score ≥ 3."),
        ("Target Attractiveness", "Average of: Visibility, Status, Threat Access, Collateral Exposure, Crime Profile, Historical Precedence."),
        ("Likelihood", "ROUND((Threat Score + Target Attractiveness + Vulnerability) / 3)."),
        ("Impact", "Average of: Disruption, Fatalities, Reputational Damage, Financial Loss, Property Damage. Life-safety override applies where Fatalities ≥ 4."),
        ("Risk Rating", "5×5 Likelihood vs Impact matrix lookup."),
        ("Control Gap", "Average of Effectiveness + Confidence: ≥ 4 = Adequate, 3–3.9 = Partial, < 3 = Insufficient."),
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    _borders(tbl)
    tbl.columns[0].width = Cm(5.0)
    tbl.columns[1].width = Cm(11.5)
    _set_cell_text(tbl.rows[0].cells[0], "Component", bold=True, size=8,
                   color=RGBColor(0xFF, 0xFF, 0xFF), fill="1B2A4A")
    _set_cell_text(tbl.rows[0].cells[1], "Formula / Description", bold=True, size=8,
                   color=RGBColor(0xFF, 0xFF, 0xFF), fill="1B2A4A")
    for i, (component, formula) in enumerate(rows, 1):
        fill = "F8FAFC" if i % 2 == 0 else None
        _set_cell_text(tbl.rows[i].cells[0], component, bold=True, size=8,
                       color=NAVY, fill=fill)
        _set_cell_text(tbl.rows[i].cells[1], formula, size=8, fill=fill)
    doc.add_paragraph()


def _build_asset_identification(doc, state):
    _add_heading(doc, "6.2  Asset Identification", 2)
    assets = [a for a in (state.get("assets", []) or []) if a.get("cat")]
    if not assets:
        _add_para(doc, "No assets recorded.", italic=True, color=GREY_MID)
        return

    _add_para(doc,
              "The following assets have been identified as being present at or associated with the "
              "site and are relevant to this security risk assessment.",
              size=10, align="justify", space_after=80)

    headers = ["#", "Asset Category", "Description"]
    tbl = doc.add_table(rows=1 + len(assets), cols=3)
    _borders(tbl)
    tbl.columns[0].width = Cm(1.0)
    tbl.columns[1].width = Cm(5.5)
    tbl.columns[2].width = Cm(10.0)
    for i, h in enumerate(headers):
        _set_cell_text(tbl.rows[0].cells[i], h, bold=True, size=8,
                       color=RGBColor(0xFF, 0xFF, 0xFF), fill="1B2A4A",
                       align="center")
    for idx, a in enumerate(assets, 1):
        fill = "F8FAFC" if idx % 2 == 0 else None
        _set_cell_text(tbl.rows[idx].cells[0], str(idx), size=8, align="center", fill=fill)
        _set_cell_text(tbl.rows[idx].cells[1], a.get("cat") or "", size=8, fill=fill)
        _set_cell_text(tbl.rows[idx].cells[2], a.get("desc") or "", size=8, fill=fill)
    doc.add_paragraph()


def _build_asset_criticality(doc, state):
    _add_heading(doc, "6.3  Asset Criticality", 2)
    assets = [a for a in (state.get("assets", []) or []) if a.get("cat")]
    if not assets:
        _add_para(doc, "No assets recorded.", italic=True, color=GREY_MID)
        return

    _add_para(doc,
              "Each asset is rated across four criteria to determine its criticality ranking. "
              "Assets scoring 4.5+ are Critical, 3.5+ High, 2.5+ Medium, 1.5+ Low, and below 1.5 Negligible.",
              size=10, align="justify", space_after=80)

    col_labels = [AC_LABELS[k] for k in AC_KEYS]
    headers = ["#", "Asset"] + col_labels + ["Score", "Ranking"]
    tbl = doc.add_table(rows=1 + len(assets), cols=len(headers))
    _borders(tbl)
    widths = [0.8, 4.5, 1.5, 1.8, 1.5, 1.8, 1.2, 1.9]
    for col, w in zip(tbl.columns, widths):
        col.width = Cm(w)
    for i, h in enumerate(headers):
        _set_cell_text(tbl.rows[0].cells[i], h, bold=True, size=8,
                       color=RGBColor(0xFF, 0xFF, 0xFF), fill="1B2A4A", align="center")

    for idx, a in enumerate(assets, 1):
        sc_obj = a.get("sc") or {}
        scores = [sc_obj.get(k) for k in AC_KEYS]
        avg_score = _avg(scores)
        rank = _asset_rank(avg_score)
        fill = "F8FAFC" if idx % 2 == 0 else None
        row = tbl.rows[idx].cells
        _set_cell_text(row[0], str(idx), size=8, align="center", fill=fill)
        _set_cell_text(row[1], a.get("cat") or "", size=8, fill=fill)
        for col_i, sc_val in enumerate(scores, 2):
            if sc_val:
                hex_fill = LEVEL_HEX.get(SCORE_LEVEL.get(int(round(sc_val))), "")
                fg = LEVEL_FG.get(SCORE_LEVEL.get(int(round(sc_val))))
                _set_cell_text(row[col_i], str(sc_val), size=8, align="center",
                                bold=True, color=fg, fill=hex_fill)
            else:
                _set_cell_text(row[col_i], "", size=8, align="center", fill=fill)
        _set_cell_text(row[6], str(avg_score) if avg_score else "", size=8,
                       align="center", bold=True, fill=fill)
        if rank:
            _set_cell_text(row[7], rank, size=8, align="center", bold=True,
                           color=CRIT_FG.get(rank), fill=CRIT_HEX.get(rank))
        else:
            _set_cell_text(row[7], "", size=8, align="center", fill=fill)
    doc.add_paragraph()


def _build_common_threat_actors(doc):
    _add_heading(doc, "6.4.1  Common Threat Actors in the UK", 3)
    _add_para(doc,
              "The following provides an overview of the principal threat actor categories relevant to "
              "security risk assessments in the United Kingdom context.",
              size=10, align="justify", space_after=80)

    _add_heading(doc, "Terrorism", 3)
    _add_para(doc,
              "The UK terrorism threat is assessed by MI5 on a five-level scale. The primary threat "
              "emanates from Islamist extremist groups, including networks inspired by or affiliated with "
              "Al-Qa'ida and Islamic State. Extreme right-wing terrorism represents a growing and significant "
              "threat, with individuals and small cells targeting minority communities and institutions. "
              "Dissident Irish Republican groups, including the New IRA and Continuity IRA, continue to "
              "operate primarily in Northern Ireland with occasional activity on the mainland. "
              "These actors employ a range of methodologies including vehicle-borne and person-borne "
              "improvised explosive devices (VBIED/PBIED), marauding terrorist firearms attacks (MTFA), "
              "stabbing attacks, and hostile reconnaissance preceding planned operations.",
              size=10, align="justify", space_after=80)

    _add_heading(doc, "Crime", 3)
    _add_para(doc,
              "Organised crime groups (OCGs) operate across the UK and pose threats including "
              "robbery, theft, fraud, cybercrime, and drug-related violence. Opportunistic criminals "
              "exploit poor physical security, inadequate surveillance, and uncontrolled access points. "
              "Insider threats — whether motivated by financial gain, coercion, or grievance — represent "
              "a particular risk where access controls or personnel vetting is insufficient. "
              "Cybercriminals increasingly target physical security systems and operational technology. "
              "The threat landscape is further shaped by lone-actor offenders, stalking and harassment, "
              "and civil disorder that may affect high-profile or publicly visible sites.",
              size=10, align="justify", space_after=80)

    doc.add_paragraph()


def _build_threat_register(doc, risks):
    _add_heading(doc, "6.4.2  Threat Register", 3)
    _add_para(doc,
              "Threats identified and assessed using the capability and intent methodology. "
              "Design Basis Threats (DBT) are those with a combined score of 3 or higher.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    filtered = [r for r in risks if r.get("n")]
    if not filtered:
        _add_para(doc, "No threats recorded.", italic=True, color=GREY_MID)
        return

    headers = ["#", "Type", "Threat", "Description", "Cap", "Int", "Score", "Level", "DBT"]
    tbl = doc.add_table(rows=1 + len(filtered), cols=len(headers))
    _borders(tbl)
    widths_cm = [0.8, 1.8, 3.5, 5.5, 0.8, 0.8, 1.0, 1.8, 1.0]
    _hdr_row(tbl, headers, widths_cm)

    for idx, r in enumerate(filtered, 1):
        fill = "F8FAFC" if idx % 2 == 0 else None
        row = tbl.rows[idx].cells
        _set_cell_text(row[0], f"T{idx}", size=8, align="center", fill=fill)
        _set_cell_text(row[1], r.get("cat") or "", size=8, align="center", fill=fill)
        _set_cell_text(row[2], r.get("n") or "", size=8, fill=fill)
        _set_cell_text(row[3], r.get("d") or "", size=8, fill=fill)
        _set_cell_text(row[4], r.get("cap") or "", size=8, align="center", fill=fill)
        _set_cell_text(row[5], r.get("int") or "", size=8, align="center", fill=fill)
        _set_cell_text(row[6], r.get("sc") or "", size=8, align="center", bold=True, fill=fill)
        lv = r.get("lv") or ""
        if lv:
            _set_cell_text(row[7], lv, size=8, align="center", bold=True,
                           color=LEVEL_FG[lv], fill=LEVEL_HEX[lv])
        else:
            _set_cell_text(row[7], "", size=8, align="center", fill=fill)
        dbt = "YES" if r.get("dbt") else ("No" if r.get("sc") is not None else "")
        _set_cell_text(row[8], dbt, size=8, align="center", bold=r.get("dbt"),
                       color=RGBColor(0xDC, 0x26, 0x26) if r.get("dbt") else None,
                       fill="FFC7CE" if r.get("dbt") else fill)
    doc.add_paragraph()


def _build_threat_asset_mapping(doc, risks, state):
    _add_heading(doc, "6.4.3  Threat-Asset Mapping", 3)
    _add_para(doc,
              "The table below maps each identified threat to the assets it targets, showing the "
              "exposure of individual assets to specific threat types.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    assets = state.get("assets", []) or []
    assets_by_id = {str(a.get("id")): a for a in assets}
    filtered = [r for r in risks if r.get("n")]

    if not filtered:
        _add_para(doc, "No threats recorded.", italic=True, color=GREY_MID)
        doc.add_paragraph()
        return

    headers = ["#", "Threat", "Type", "Targeted Assets"]
    tbl = doc.add_table(rows=1 + len(filtered), cols=len(headers))
    _borders(tbl)
    widths_cm = [0.8, 4.0, 1.8, 10.0]
    _hdr_row(tbl, headers, widths_cm)

    for idx, r in enumerate(filtered, 1):
        fill = "F8FAFC" if idx % 2 == 0 else None
        row = tbl.rows[idx].cells
        _set_cell_text(row[0], f"T{idx}", size=8, align="center", fill=fill)
        _set_cell_text(row[1], r.get("n") or "", size=8, fill=fill)
        _set_cell_text(row[2], r.get("cat") or "", size=8, align="center", fill=fill)
        target_names = []
        for tid in (r.get("targets") or []):
            a = assets_by_id.get(str(tid))
            if a and a.get("cat"):
                label = a["cat"]
                if a.get("desc"):
                    label += f" — {a['desc']}"
                target_names.append(label)
        _set_cell_text(row[3], "; ".join(target_names) if target_names else "—",
                       size=8, fill=fill)
    doc.add_paragraph()


def _build_attractiveness_vulnerability(doc, risks, state):
    _add_heading(doc, "6.5  Attractiveness, Vulnerability & Likelihood", 2)
    _add_para(doc,
              "Target attractiveness is scored across six criteria. Combined with the threat score and "
              "site vulnerability, this determines the overall likelihood rating.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    vs = state.get("vs", {}) or {}
    filtered = [r for r in risks if r.get("n")]
    if not filtered:
        _add_para(doc, "No vulnerability data recorded.", italic=True, color=GREY_MID)
        return

    tac_short = ["Vis", "Stat", "Acc", "Coll", "Crime", "Hist"]
    headers = ["#", "Threat", "T", "Vis", "Stat", "Acc", "Coll", "Crime", "Hist", "TA", "Vuln", "L"]
    tbl = doc.add_table(rows=1 + len(filtered), cols=len(headers))
    _borders(tbl)
    widths_cm = [0.8, 4.5, 0.8, 0.8, 0.8, 0.8, 0.8, 1.0, 0.8, 0.8, 0.8, 0.8]
    _hdr_row(tbl, headers, widths_cm)

    for idx, r in enumerate(filtered, 1):
        fill = "F8FAFC" if idx % 2 == 0 else None
        v = vs.get(str(r.get("id")), {}) or {}
        row = tbl.rows[idx].cells
        _set_cell_text(row[0], f"T{idx}", size=8, align="center", fill=fill)
        _set_cell_text(row[1], r.get("n") or "", size=8, fill=fill)

        def _score_cell(cell, val):
            if val:
                lv = SCORE_LEVEL.get(int(round(val)))
                _set_cell_text(cell, str(val), size=8, align="center", bold=True,
                               color=LEVEL_FG.get(lv), fill=LEVEL_HEX.get(lv))
            else:
                _set_cell_text(cell, "", size=8, align="center", fill=fill)

        _score_cell(row[2], r.get("sc"))
        for col_i, tac_key in enumerate(TAC, 3):
            _score_cell(row[col_i], v.get(tac_key))
        _score_cell(row[9], r.get("ta"))
        _score_cell(row[10], r.get("vu"))
        _score_cell(row[11], r.get("lik"))
    doc.add_paragraph()


def _build_points_of_interest(doc, state):
    _add_heading(doc, "6.6  Points of Interest", 2)
    pois = state.get("pois", []) or []
    threats = state.get("threats", []) or []
    threats_by_id = {str(t.get("id")): t for t in threats}

    if not pois:
        _add_para(doc, "No Points of Interest recorded.", italic=True, color=GREY_MID)
        return

    _add_para(doc,
              "Points of Interest (POIs) are locations or features within or adjacent to the site "
              "that present specific security considerations.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    headers = ["#", "POI Name", "Location / Description", "Applicable Threats"]
    tbl = doc.add_table(rows=1 + len(pois), cols=len(headers))
    _borders(tbl)
    widths_cm = [0.8, 3.5, 6.0, 6.2]
    _hdr_row(tbl, headers, widths_cm)

    for idx, poi in enumerate(pois, 1):
        fill = "F8FAFC" if idx % 2 == 0 else None
        threat_names = []
        for tid in (poi.get("threats") or []):
            t = threats_by_id.get(str(tid))
            if t and t.get("n"):
                threat_names.append(t["n"])
        row = tbl.rows[idx].cells
        _set_cell_text(row[0], f"POI-{idx}", size=8, align="center", fill=fill)
        _set_cell_text(row[1], poi.get("name") or "", size=8, fill=fill)
        _set_cell_text(row[2], poi.get("desc") or "", size=8, fill=fill)
        _set_cell_text(row[3], ", ".join(threat_names) if threat_names else "—",
                       size=8, fill=fill)
    doc.add_paragraph()


def _build_impact_assessment(doc, risks, state):
    _add_heading(doc, "6.7  Impact Assessment", 2)
    _add_para(doc,
              "Impact is assessed across five consequence categories. A life-safety override applies "
              "where Fatalities score ≥ 4 and the average impact is below 4.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    is_ = state.get("is_", {}) or {}
    filtered = [r for r in risks if r.get("n")]
    if not filtered:
        _add_para(doc, "No impact data recorded.", italic=True, color=GREY_MID)
        return

    ic_short = ["Disruption", "Fatalities", "Reputation", "Financial", "Property"]
    headers = ["#", "Threat"] + ic_short + ["Impact", "Override?"]
    tbl = doc.add_table(rows=1 + len(filtered), cols=len(headers))
    _borders(tbl)
    widths_cm = [0.8, 4.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.2, 1.5]
    _hdr_row(tbl, headers, widths_cm)

    for idx, r in enumerate(filtered, 1):
        fill = "F8FAFC" if idx % 2 == 0 else None
        im = is_.get(str(r.get("id")), {}) or {}
        row = tbl.rows[idx].cells
        _set_cell_text(row[0], f"T{idx}", size=8, align="center", fill=fill)
        _set_cell_text(row[1], r.get("n") or "", size=8, fill=fill)

        def _score_cell(cell, val):
            if val:
                lv = SCORE_LEVEL.get(int(round(val)))
                _set_cell_text(cell, str(val), size=8, align="center", bold=True,
                               color=LEVEL_FG.get(lv), fill=LEVEL_HEX.get(lv))
            else:
                _set_cell_text(cell, "", size=8, align="center", fill=fill)

        for col_i, ic_key in enumerate(IC, 2):
            _score_cell(row[col_i], im.get(ic_key))
        _score_cell(row[7], r.get("imp"))
        override = "YES" if r.get("lifeSafetyOverride") else ""
        _set_cell_text(row[8], override, size=8, align="center", bold=bool(override),
                       color=RGBColor(0xDC, 0x26, 0x26) if override else None,
                       fill="FFC7CE" if override else fill)
    doc.add_paragraph()


def _build_existing_controls(doc, risks, state):
    _add_heading(doc, "6.8  Existing Controls", 2)
    _add_para(doc,
              "Existing security controls are evaluated for effectiveness and confidence. "
              "The gap assessment indicates whether current controls are Adequate, Partial, or Insufficient.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    cd = state.get("cd", {}) or {}
    filtered = [r for r in risks if r.get("n")]
    if not filtered:
        _add_para(doc, "No control data recorded.", italic=True, color=GREY_MID)
        return

    headers = ["T#", "Threat", "Controls Description", "Eff", "Conf", "Score", "Gap"]
    tbl = doc.add_table(rows=1 + len(filtered), cols=len(headers))
    _borders(tbl)
    widths_cm = [0.8, 3.5, 7.5, 0.8, 0.8, 0.9, 2.2]
    _hdr_row(tbl, headers, widths_cm)

    for idx, r in enumerate(filtered, 1):
        fill = "F8FAFC" if idx % 2 == 0 else None
        c = cd.get(str(r.get("id")), {}) or {}
        eff, conf = c.get("eff"), c.get("conf")
        gap_score = round((eff + conf) / 2 * 10) / 10 if (eff and conf) else None
        gap = ("Adequate" if gap_score >= 4 else
               "Partial" if gap_score >= 3 else
               "Insufficient" if gap_score else "")
        row = tbl.rows[idx].cells
        _set_cell_text(row[0], f"T{idx}", size=8, align="center", fill=fill)
        _set_cell_text(row[1], r.get("n") or "", size=8, fill=fill)
        _set_cell_text(row[2], c.get("desc") or "", size=8, fill=fill)

        def _score_cell(cell, val):
            if val:
                lv = SCORE_LEVEL.get(int(round(val)))
                _set_cell_text(cell, str(val), size=8, align="center", bold=True,
                               color=LEVEL_FG.get(lv), fill=LEVEL_HEX.get(lv))
            else:
                _set_cell_text(cell, "", size=8, align="center", fill=fill)

        _score_cell(row[3], eff)
        _score_cell(row[4], conf)
        _set_cell_text(row[5], str(gap_score) if gap_score else "", size=8,
                       align="center", bold=True, fill=fill)
        if gap:
            _set_cell_text(row[6], gap, size=8, align="center", bold=True,
                           color=GAP_FG.get(gap), fill=GAP_HEX.get(gap))
        else:
            _set_cell_text(row[6], "", size=8, align="center", fill=fill)
    doc.add_paragraph()


def _build_risk_register(doc, risks):
    _add_heading(doc, "6.9  Risk Register (Pre Mitigation)", 2)
    _add_para(doc,
              "Pre-treatment risk ratings derived from the 5×5 Likelihood vs Impact matrix. "
              "Threat Score (T), Target Attractiveness (TA), and Vulnerability (V) combine to produce the Likelihood (L) rating.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    filtered = [r for r in risks if r.get("n")]
    if not filtered:
        _add_para(doc, "No risk data recorded.", italic=True, color=GREY_MID)
        return

    headers = ["R#", "Type", "Threat", "T", "TA", "V", "L", "I", "Rating", "Treatment"]
    tbl = doc.add_table(rows=1 + len(filtered), cols=len(headers))
    _borders(tbl)
    widths_cm = [0.8, 1.8, 4.0, 0.7, 0.7, 0.7, 0.7, 0.7, 1.8, 2.1]
    _hdr_row(tbl, headers, widths_cm)

    for idx, r in enumerate(filtered, 1):
        fill = "F8FAFC" if idx % 2 == 0 else None
        row = tbl.rows[idx].cells
        _set_cell_text(row[0], f"R{idx}", size=8, align="center", fill=fill)
        _set_cell_text(row[1], r.get("cat") or "", size=8, align="center", fill=fill)
        _set_cell_text(row[2], r.get("n") or "", size=8, fill=fill)

        def _score_cell(cell, val):
            if val:
                lv = SCORE_LEVEL.get(int(round(val)))
                _set_cell_text(cell, str(val), size=8, align="center", bold=True,
                               color=LEVEL_FG.get(lv), fill=LEVEL_HEX.get(lv))
            else:
                _set_cell_text(cell, "", size=8, align="center", fill=fill)

        _score_cell(row[3], r.get("sc"))
        _score_cell(row[4], r.get("ta"))
        _score_cell(row[5], r.get("vu"))
        _score_cell(row[6], r.get("lik"))
        _score_cell(row[7], r.get("imp"))
        rat = r.get("rat") or ""
        if rat:
            _set_cell_text(row[8], rat, size=8, align="center", bold=True,
                           color=LEVEL_FG[rat], fill=LEVEL_HEX[rat])
        else:
            _set_cell_text(row[8], "", size=8, align="center", fill=fill)
        _set_cell_text(row[9], r.get("treat") or "", size=8, align="center", fill=fill)

    doc.add_paragraph()
    _add_heading(doc, "Pre-Mitigation Risk Matrix", 3)
    _try_picture(doc, lambda: chart_risk_matrix(risks, residual=False,
                                                title="Likelihood × Impact (Pre-Mitigation)"),
                 Cm(15))
    _page_break(doc)


def _build_security_treatments(doc, risks, state):
    _add_heading(doc, "6.10  Security Treatments", 2)
    _add_para(doc,
              "Treatment strategies follow the TAAR framework: Treat, Avoid, Accept, Refer. "
              "Planned measures for headline risks (High or Very High) are listed per affected asset. "
              "The residual risk position following treatment is also presented.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    treatment_measures = state.get("treatmentMeasures", {}) or {}
    assets = state.get("assets", []) or []
    assets_by_id = {str(a.get("id")): a for a in assets}

    treat_counts = {k: 0 for k in ("Reduce", "Transfer", "Avoid", "Accept")}
    for r in risks:
        t = r.get("treat")
        if t in treat_counts:
            treat_counts[t] += 1
    _add_para(doc,
              "Treatment strategy split: "
              + ", ".join(f"{v} {k}" for k, v in treat_counts.items() if v),
              size=10, space_after=100)

    headline = [r for r in risks if r.get("rat") in ("Very High", "High")]
    if headline:
        _add_heading(doc, "Treatment Measures by Risk", 3)
        for r in headline:
            tid = str(r.get("id"))
            m = treatment_measures.get(tid, {}) or {}
            any_measure = any(any(v for v in (arr or []) if v) for arr in m.values())
            if not any_measure:
                continue

            _add_heading(doc, f"{r.get('n')}  —  {r.get('rat')} ({r.get('treat') or 'Reduce'})", 3)

            rows_data = []
            for target_id in (r.get("targets") or []):
                a = assets_by_id.get(str(target_id))
                arr = [v for v in (m.get(str(target_id)) or []) if v]
                if a and a.get("cat") and arr:
                    rows_data.append((a, arr))
            if not rows_data:
                continue

            tbl = doc.add_table(rows=1 + len(rows_data), cols=2)
            _borders(tbl)
            tbl.columns[0].width = Cm(5.5)
            tbl.columns[1].width = Cm(11.0)
            _set_cell_text(tbl.rows[0].cells[0], "Asset / Location", bold=True, size=9,
                           color=RGBColor(0xFF, 0xFF, 0xFF), fill="1B2A4A")
            _set_cell_text(tbl.rows[0].cells[1], "Treatment Measures", bold=True, size=9,
                           color=RGBColor(0xFF, 0xFF, 0xFF), fill="1B2A4A")
            for i, (a, arr) in enumerate(rows_data, 1):
                c0 = tbl.rows[i].cells[0]
                c1 = tbl.rows[i].cells[1]
                c0.text = ""
                p = c0.paragraphs[0]
                rn = p.add_run(a.get("cat") or "")
                rn.font.bold = True; rn.font.size = Pt(9); rn.font.color.rgb = NAVY
                if a.get("desc"):
                    p2 = c0.add_paragraph()
                    r2 = p2.add_run(a.get("desc"))
                    r2.font.size = Pt(8); r2.font.color.rgb = GREY_DARK
                c1.text = ""
                for j, measure in enumerate(arr, 1):
                    p = c1.paragraphs[0] if j == 1 else c1.add_paragraph()
                    run = p.add_run(f"{j}. {measure}")
                    run.font.size = Pt(8); run.font.color.rgb = GREY_DARK
                    _set_para_spacing(p, after=30)
            doc.add_paragraph()

    # Residual risk table
    _add_heading(doc, "Residual Risk", 3)
    _add_para(doc,
              "Risk rating following application of proposed treatment measures.",
              size=9, italic=True, color=GREY_MID, space_after=80)

    filtered = [r for r in risks if r.get("n") and (r.get("rat") or r.get("resRat"))]
    if filtered:
        headers = ["#", "Threat", "Pre L", "Pre I", "Pre-Rating",
                   "Treatment", "Post L", "Post I", "Residual", "Accept?"]
        tbl = doc.add_table(rows=1 + len(filtered), cols=len(headers))
        _borders(tbl)
        widths = [0.8, 4.0, 0.9, 0.9, 1.8, 1.8, 0.9, 0.9, 1.8, 1.2]
        _hdr_row(tbl, headers, widths)

        for idx, r in enumerate(filtered, 1):
            fill = "F8FAFC" if idx % 2 == 0 else None
            row = tbl.rows[idx].cells
            _set_cell_text(row[0], f"R{idx}", size=8, align="center", fill=fill)
            _set_cell_text(row[1], r.get("n") or "", size=8, fill=fill)

            def _score_cell(cell, val):
                if val:
                    lv = SCORE_LEVEL.get(int(round(val)))
                    _set_cell_text(cell, str(val), size=8, align="center", bold=True,
                                   color=LEVEL_FG.get(lv), fill=LEVEL_HEX.get(lv))
                else:
                    _set_cell_text(cell, "", size=8, align="center", fill=fill)

            _score_cell(row[2], r.get("lik"))
            _score_cell(row[3], r.get("imp"))
            rat = r.get("rat") or ""
            if rat:
                _set_cell_text(row[4], rat, size=8, align="center", bold=True,
                               color=LEVEL_FG[rat], fill=LEVEL_HEX[rat])
            else:
                _set_cell_text(row[4], "", size=8, align="center", fill=fill)
            _set_cell_text(row[5], r.get("treat") or "", size=8, align="center", fill=fill)
            _score_cell(row[6], r.get("resL"))
            _score_cell(row[7], r.get("resI"))
            res = r.get("resRat") or ""
            if res:
                _set_cell_text(row[8], res, size=8, align="center", bold=True,
                               color=LEVEL_FG[res], fill=LEVEL_HEX[res])
            else:
                _set_cell_text(row[8], "", size=8, align="center", fill=fill)
            _set_cell_text(row[9], r.get("acc") or "", size=8, align="center", fill=fill)

        doc.add_paragraph()
        _add_heading(doc, "Post-Mitigation Risk Matrix", 3)
        _try_picture(doc, lambda: chart_risk_matrix(risks, residual=True,
                                                    title="Likelihood × Impact (Residual)"),
                     Cm(15))
    else:
        _add_para(doc, "Residual risk scoring not yet captured.", italic=True, color=GREY_MID)


# --- Top-level build ---

def build_docx(state):
    doc = Document()

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = GREY_DARK

    risks = compute_risks(state)
    ctx = state.get("ctx", {}) or {}
    ctx["_v"] = state.get("_v") or ctx.get("_v") or ""
    pois = state.get("pois", []) or []

    _build_cover(doc, ctx, risks)
    _build_exec_summary(doc, ctx, risks, pois)

    # Section 6
    _build_s6_header(doc)
    _build_methodology(doc)
    _page_break(doc)

    _build_asset_identification(doc, state)
    _build_asset_criticality(doc, state)
    _page_break(doc)

    _add_heading(doc, "6.4  Threats", 2)
    _build_common_threat_actors(doc)
    _build_threat_register(doc, risks)
    _page_break(doc)
    _build_threat_asset_mapping(doc, risks, state)
    _page_break(doc)

    _build_attractiveness_vulnerability(doc, risks, state)
    _page_break(doc)

    _build_points_of_interest(doc, state)
    _page_break(doc)

    _build_impact_assessment(doc, risks, state)
    _page_break(doc)

    _build_existing_controls(doc, risks, state)
    _page_break(doc)

    _build_risk_register(doc, risks)

    _build_security_treatments(doc, risks, state)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _safe_filename(name):
    s = re.sub(r"[^A-Za-z0-9_-]+", "_", name or "Assessment")
    return s.strip("_") or "Assessment"


# --- Vercel handler ---

def _unauthorized(self):
    msg = json.dumps({"error": "unauthorized"}).encode("utf-8")
    self.send_response(401)
    self.send_header("Content-Type", "application/json")
    self.send_header("Content-Length", str(len(msg)))
    self.end_headers()
    self.wfile.write(msg)


class handler(BaseHTTPRequestHandler):
    def _authorized(self):
        expected = os.environ.get("APP_PASSWORD") or ""
        if not expected:
            return False
        provided = self.headers.get("x-api-key") or ""
        return provided == expected

    def do_POST(self):
        if not self._authorized():
            _unauthorized(self)
            return
        try:
            length = int(self.headers.get("content-length", 0))
            body = self.rfile.read(length) if length else b"{}"
            state = json.loads(body.decode("utf-8"))
            docx_bytes = build_docx(state)
            filename = _safe_filename((state.get("ctx") or {}).get("project")) + "_SRA.docx"
            self.send_response(200)
            self.send_header("Content-Type",
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
            self.send_header("Content-Length", str(len(docx_bytes)))
            self.end_headers()
            self.wfile.write(docx_bytes)
        except Exception as e:
            msg = json.dumps({"error": str(e)}).encode("utf-8")
            self.send_response(500)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(msg)))
            self.end_headers()
            self.wfile.write(msg)

    def do_GET(self):
        msg = json.dumps({
            "ok": True,
            "endpoint": "/api/report",
            "method": "POST",
            "python": os.sys.version.split()[0],
        }).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(msg)))
        self.end_headers()
        self.wfile.write(msg)

    def do_OPTIONS(self):
        self.send_response(204)
        self.end_headers()
